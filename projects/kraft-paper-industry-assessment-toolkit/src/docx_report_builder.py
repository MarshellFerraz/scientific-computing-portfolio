from __future__ import annotations

from pathlib import Path
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)


def build_docx_report(evaluation: Dict[str, Any], output_path: str | Path) -> None:
    """Build a simple editable DOCX feedback report from evaluation JSON.

    This is intentionally minimal. Institutions can replace it with their own
    approved DOCX template.
    """
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Rubrica de Correção e Devolutiva Individual")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(31, 78, 121)

    document.add_paragraph("Prova - Papel e Celulose", style=None).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Planta integrada de celulose e papel baseada no processo Kraft").alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(document, "Quadro de notas por tópico", 1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Critério"
    hdr[2].text = "Máximo"
    hdr[3].text = "Nota sugerida"

    for result in evaluation["results"]:
        row = table.add_row().cells
        row[0].text = str(result["item"])
        row[1].text = result["name"]
        row[2].text = str(result["max_score"])
        row[3].text = str(result["suggested_score"])

    add_heading(document, "Ficha de devolutiva por tópico", 1)
    for result in evaluation["results"]:
        add_heading(document, f"{result['item']}. {result['name']} - {result['suggested_score']} / {result['max_score']}", 2)
        document.add_paragraph(result["feedback"])

    document.add_paragraph("Observação: pontuação preliminar; validar tecnicamente antes da emissão final.")
    document.save(output_path)
